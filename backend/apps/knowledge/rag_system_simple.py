"""
简化的RAG系统实现，避免sklearn依赖问题
"""
import os
import json
import logging
import hashlib
import asyncio
from typing import List, Dict, Optional, Tuple, Any
from datetime import datetime
import uuid

# 文档处理
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from markdown import markdown
    from bs4 import BeautifulSoup
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

# 向量化和检索
import numpy as np

# 文本处理
try:
    import jieba
    HAS_JIEBA = True
except ImportError:
    HAS_JIEBA = False

import re

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器"""
    
    @staticmethod
    def process_file(file_path: str) -> Tuple[str, Dict]:
        """处理单个文件"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.txt':
            return DocumentProcessor._process_txt(file_path)
        elif file_ext == '.md':
            return DocumentProcessor._process_markdown(file_path)
        elif file_ext == '.pdf' and HAS_PDF:
            return DocumentProcessor._process_pdf(file_path)
        elif file_ext in ['.docx', '.doc'] and HAS_DOCX:
            return DocumentProcessor._process_docx(file_path)
        elif file_ext == '.html':
            return DocumentProcessor._process_html(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    @staticmethod
    def _process_txt(file_path: str) -> Tuple[str, Dict]:
        """处理TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                content = f.read()
        
        metadata = {
            'source': file_path,
            'type': 'txt',
            'size': len(content)
        }
        
        return content, metadata
    
    @staticmethod
    def _process_markdown(file_path: str) -> Tuple[str, Dict]:
        """处理Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取纯文本（移除markdown语法）
        if HAS_MARKDOWN:
            html = markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
        else:
            # 简单的markdown语法移除
            text = re.sub(r'[#*`_\[\]()]', '', content)
        
        metadata = {
            'source': file_path,
            'type': 'markdown',
            'size': len(text)
        }
        
        return text, metadata
    
    @staticmethod
    def _process_pdf(file_path: str) -> Tuple[str, Dict]:
        """处理PDF文件"""
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        metadata = {
            'source': file_path,
            'type': 'pdf',
            'pages': len(reader.pages),
            'size': len(text)
        }
        
        return text, metadata
    
    @staticmethod
    def _process_docx(file_path: str) -> Tuple[str, Dict]:
        """处理DOCX文件"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        metadata = {
            'source': file_path,
            'type': 'docx',
            'paragraphs': len(doc.paragraphs),
            'size': len(text)
        }
        
        return text, metadata
    
    @staticmethod
    def _process_html(file_path: str) -> Tuple[str, Dict]:
        """处理HTML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        if HAS_MARKDOWN:
            soup = BeautifulSoup(content, 'html.parser')
            text = soup.get_text()
        else:
            # 简单的HTML标签移除
            text = re.sub(r'<[^>]+>', '', content)
        
        metadata = {
            'source': file_path,
            'type': 'html',
            'size': len(text)
        }
        
        return text, metadata


class TextSplitter:
    """文本分块器"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """分割文本为块"""
        # 简单的文本分割
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # 尝试在句号处分割
            if end < len(text):
                last_period = chunk.rfind('。')
                if last_period > self.chunk_size // 2:
                    chunk = chunk[:last_period + 1]
                    end = start + last_period + 1
            
            chunk_metadata = {
                'chunk_index': len(chunks),
                'chunk_size': len(chunk),
                'chunk_word_count': len(chunk.split()),
                **(metadata or {})
            }
            
            chunks.append({
                'content': chunk.strip(),
                'metadata': chunk_metadata
            })
            
            start = end - self.chunk_overlap
        
        return chunks


class SimpleEmbedding:
    """简单的嵌入模型实现"""
    
    def __init__(self, vector_size=300):
        self.vector_size = vector_size
        self.is_fitted = True  # 简化版本，不需要训练
        # 预定义词汇表和哈希函数
        self.vocab = {}
        self.vocab_size = 1000  # 限制词汇表大小
    
    def _get_vocab(self, texts):
        """建立词汇表"""
        vocab = set()
        for text in texts:
            # 简单的中文分词（按字符）
            chars = list(text)
            vocab.update(chars)
        
        # 限制词汇表大小，选择最常见的字符
        if len(vocab) > self.vocab_size:
            from collections import Counter
            all_chars = []
            for text in texts:
                all_chars.extend(list(text))
            char_counts = Counter(all_chars)
            vocab = [char for char, count in char_counts.most_common(self.vocab_size)]
        else:
            vocab = list(vocab)
        
        return vocab
    
    def encode(self, texts: List[str]) -> np.ndarray:
        """编码文本为向量"""
        if not texts:
            return np.array([])
        
        # 如果是第一次编码，建立词汇表
        if not self.vocab:
            vocab_list = self._get_vocab(texts)
            self.vocab = {char: i for i, char in enumerate(vocab_list)}
        
        vectors = []
        for text in texts:
            # 创建固定大小的向量
            vector = np.zeros(self.vector_size)
            
            # 字符级别的简单编码
            chars = list(text)
            for i, char in enumerate(chars):
                if char in self.vocab:
                    idx = self.vocab[char]
                    # 使用哈希函数映射到固定大小的向量
                    pos = idx % self.vector_size
                    vector[pos] += 1
            
            # 归一化
            norm = np.linalg.norm(vector)
            if norm > 0:
                vector = vector / norm
            
            vectors.append(vector)
        
        return np.array(vectors, dtype=float)


class VectorStore:
    """向量存储器"""
    
    def __init__(self, embedding_model=None):
        self.embedding_model = embedding_model or SimpleEmbedding()
        self.chunks = []
        self.vectors = None
        self.metadata = []
    
    def add_documents(self, chunks: List[Dict]):
        """添加文档块并持久化到数据库"""
        from django.db import transaction
        from apps.knowledge.models import DocumentChunk, Document
        
        with transaction.atomic():
            for chunk in chunks:
                self.chunks.append(chunk['content'])
                self.metadata.append(chunk['metadata'])
                
                # 持久化到数据库
                if 'document_id' in chunk['metadata']:
                    try:
                        document = Document.objects.get(id=chunk['metadata']['document_id'])
                        DocumentChunk.objects.create(
                            document=document,
                            chunk_index=chunk['metadata'].get('chunk_index', 0),
                            content=chunk['content'],
                            embedding=None,  # 向量将在_update_vectors中更新
                            metadata=chunk['metadata']
                        )
                    except Document.DoesNotExist:
                        logger.warning(f"Document with ID {chunk['metadata']['document_id']} not found")
        
        # 重新计算向量
        self._update_vectors()
    
    def _update_vectors(self):
        """更新向量并持久化到数据库"""
        if self.chunks:
            self.vectors = self.embedding_model.encode(self.chunks)
            
            # 更新数据库中的向量 - 处理异步环境
            from apps.knowledge.models import DocumentChunk
            import asyncio
            import threading
            import queue
            
            # 检查是否在异步环境中
            try:
                loop = asyncio.get_event_loop()
                if loop.is_running():
                    # 在异步环境中，使用线程来执行同步数据库操作
                    def sync_update_embeddings():
                        for i, vector in enumerate(self.vectors):
                            if i < len(self.metadata) and 'document_id' in self.metadata[i]:
                                try:
                                    chunk = DocumentChunk.objects.filter(
                                        document_id=self.metadata[i]['document_id'],
                                        chunk_index=self.metadata[i].get('chunk_index', 0)
                                    ).first()
                                    if chunk:
                                        chunk.embedding = vector.tolist()
                                        chunk.save()
                                except Exception as e:
                                    logger.warning(f"Failed to update embedding for chunk {i}: {e}")
                    
                    # 在独立线程中执行数据库更新
                    thread = threading.Thread(target=sync_update_embeddings)
                    thread.start()
                    thread.join()
                    return
                    
            except RuntimeError:
                # 没有事件循环，可以进行同步操作
                pass
            
            # 同步环境中的正常更新
            for i, vector in enumerate(self.vectors):
                if i < len(self.metadata) and 'document_id' in self.metadata[i]:
                    try:
                        chunk = DocumentChunk.objects.filter(
                            document_id=self.metadata[i]['document_id'],
                            chunk_index=self.metadata[i].get('chunk_index', 0)
                        ).first()
                        if chunk:
                            chunk.embedding = vector.tolist()
                            chunk.save()
                    except Exception as e:
                        logger.warning(f"Failed to update embedding for chunk {i}: {e}")
    
    def similarity_search(self, query: str, top_k: int = 5, threshold: float = 0.1) -> List[Dict]:
        """相似度搜索"""
        if not self.chunks or self.vectors is None:
            return []
        
        # 编码查询
        query_vector = self.embedding_model.encode([query])
        
        # 简单的余弦相似度计算
        def cosine_sim(a, b):
            dot_product = np.dot(a, b)
            norm_a = np.linalg.norm(a)
            norm_b = np.linalg.norm(b)
            if norm_a == 0 or norm_b == 0:
                return 0
            return dot_product / (norm_a * norm_b)
        
        similarities = []
        for vector in self.vectors:
            sim = cosine_sim(query_vector[0], vector)
            similarities.append(sim)
        similarities = np.array(similarities)
        
        # 获取top_k结果
        top_indices = np.argsort(similarities)[::-1][:top_k]
        
        results = []
        for idx in top_indices:
            score = similarities[idx]
            if score >= threshold:
                results.append({
                    'content': self.chunks[idx],
                    'score': float(score),
                    'metadata': self.metadata[idx],
                    'index': int(idx)
                })
        
        return results


class LLMInterface:
    """大语言模型接口"""
    
    def __init__(self, model_config):
        """初始化LLM接口
        
        Args:
            model_config: 可以是字典或ModelConfig对象
        """
        if hasattr(model_config, 'model_type'):
            # 如果是ModelConfig对象，转换为字典
            self.model_config = {
                'model_type': model_config.model_type,
                'model_name': model_config.model_name,
                'api_key': model_config.api_key,
                'api_base_url': model_config.api_base_url,
                'max_tokens': model_config.max_tokens,
                'temperature': model_config.temperature,
            }
        else:
            # 如果是字典，直接使用
            self.model_config = model_config
        
        self.model_type = self.model_config.get('model_type', 'mock')
        self.model_name = self.model_config.get('model_name', 'mock')
    
    async def generate(self, prompt: str, context: str = "") -> str:
        """生成回答"""
        if self.model_type == 'mock' or self.model_name == 'mock':
            return f"基于提供的上下文信息：{context[:100]}...\n\n对于问题「{prompt}」，这是一个模拟回答。请配置真实的大语言模型以获得准确回答。"
        
        # 真实API调用
        if self.model_type == 'api':
            try:
                return await self._call_real_api(prompt, context)
            except Exception as e:
                logger.error(f"API调用失败: {e}")
                return f"抱歉，调用AI模型时出现错误：{str(e)}"
        
        return "请配置大语言模型"
    
    async def _call_real_api(self, prompt: str, context: str = "") -> str:
        """调用真实的API"""
        import aiohttp
        import json
        
        api_key = self.model_config.get('api_key')
        api_base_url = self.model_config.get('api_base_url')
        model_name = self.model_config.get('model_name')
        
        if not api_key:
            return "错误：未配置API密钥"
        
        # 直接使用传入的prompt，不再重新构建
        # 因为在ask_question中已经构建了完整的提示
        full_prompt = prompt
        
        # 记录发送给API的提示内容（截取前200字符）
        logger.info(f"发送给API的提示预览: {full_prompt[:200]}...")
        
        try:
            # 支持不同的API格式
            if 'gemini' in model_name.lower() or 'google' in api_base_url.lower():
                return await self._call_gemini_api(full_prompt, api_key, api_base_url)
            elif 'openai' in api_base_url.lower():
                return await self._call_openai_api(full_prompt, api_key, api_base_url, model_name)
            else:
                return await self._call_generic_api(full_prompt, api_key, api_base_url, model_name)
        except Exception as e:
            logger.error(f"API调用异常: {e}")
            return f"API调用失败：{str(e)}"
    
    async def _call_gemini_api(self, prompt: str, api_key: str, api_base_url: str) -> str:
        """调用Gemini API - 使用同步方式确保稳定性"""
        import requests
        import json
        import asyncio
        
        # 在线程池中执行同步请求，避免阻塞事件循环
        def sync_request():
            # Gemini API URL格式 - 使用v1beta
            if not api_base_url.endswith('/'):
                api_base_url_fixed = api_base_url + '/'
            else:
                api_base_url_fixed = api_base_url
            
            # 使用配置中的模型名称
            model_name = self.model_config.get('model_name', 'gemini-pro')
            url = f"{api_base_url_fixed}v1beta/models/{model_name}:generateContent"
            
            # 正确的请求头格式 - 使用x-goog-api-key
            headers = {
                'Content-Type': 'application/json',
                'x-goog-api-key': api_key
            }
            
            # 正确的请求体格式
            data = {
                "contents": [
                    {
                        "role": "user",
                        "parts": [
                            {
                                "text": prompt
                            }
                        ]
                    }
                ],
                "generationConfig": {
                    "temperature": self.model_config.get('temperature', 0.7),
                    "maxOutputTokens": self.model_config.get('max_tokens', 4096),
                }
            }
            
            logger.info(f"Gemini API URL: {url}")
            logger.info(f"Gemini API Headers: {headers}")
            
            try:
                response = requests.post(url, headers=headers, json=data, timeout=30)
                logger.info(f"Gemini API Response Status: {response.status_code}")
                
                if response.status_code == 200:
                    result = response.json()
                    candidates = result.get('candidates', [])
                    if candidates and candidates[0].get('content'):
                        parts = candidates[0]['content'].get('parts', [])
                        if parts:
                            answer = parts[0].get('text', '未获得有效回复')
                            logger.info(f"Gemini API Success: {answer[:100]}...")
                            return answer
                    logger.warning("Gemini API响应格式不正确")
                    return '未获得有效回复'
                else:
                    logger.error(f"Gemini API请求失败 ({response.status_code}): {response.text}")
                    raise Exception(f"API请求失败 ({response.status_code}): {response.text}")
                    
            except Exception as e:
                logger.error(f"Gemini API调用异常: {e}")
                raise e
        
        # 在线程池中执行同步请求
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(None, sync_request)
        return result
    
    async def _call_openai_api(self, prompt: str, api_key: str, api_base_url: str, model_name: str) -> str:
        """调用OpenAI API"""
        import aiohttp
        import json
        
        url = f"{api_base_url}/chat/completions"
        
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_key}'
        }
        
        data = {
            "model": model_name,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": self.model_config.get('temperature', 0.7),
            "max_tokens": self.model_config.get('max_tokens', 4096),
        }
        
        timeout = aiohttp.ClientTimeout(total=30)
        async with aiohttp.ClientSession(timeout=timeout) as session:
            async with session.post(url, headers=headers, json=data) as response:
                if response.status == 200:
                    result = await response.json()
                    choices = result.get('choices', [])
                    if choices:
                        return choices[0]['message']['content']
                    return '未获得有效回复'
                else:
                    error_text = await response.text()
                    raise Exception(f"API请求失败 ({response.status}): {error_text}")
    
    async def _call_generic_api(self, prompt: str, api_key: str, api_base_url: str, model_name: str) -> str:
        """调用通用API"""
        import aiohttp
        import json
        
        # 尝试OpenAI格式
        try:
            return await self._call_openai_api(prompt, api_key, api_base_url, model_name)
        except Exception as e:
            logger.error(f"通用API调用失败: {e}")
            return f"API调用失败：{str(e)}"
    
    async def generate_response(self, prompt: str, context: str = "") -> Dict:
        """生成回答并返回详细信息"""
        import time
        start_time = time.time()
        
        logger.info(f"LLM 开始生成回答，模型: {self.model_name}")
        logger.info(f"传入的提示词长度: {len(prompt)} 字符")
        logger.info(f"提示词前200字符: {prompt[:200]}...")
        
        try:
            # 注意：这里不传递context，因为prompt已经包含了完整的提示词
            answer = await self.generate(prompt, "")
            response_time = round((time.time() - start_time), 3)  # 保持为秒，保留3位小数
            
            logger.info(f"LLM 回答生成成功，耗时: {response_time}秒")
            logger.info(f"回答预览: {answer[:200]}...")
            
            return {
                'answer': answer,
                'response_time': response_time,
                'model_used': f"{self.model_name}",
                'success': True
            }
        except Exception as e:
            response_time = round((time.time() - start_time), 3)  # 保持为秒，保留3位小数
            logger.error(f"LLM 回答生成失败: {str(e)}")
            return {
                'answer': f"生成回答时出错: {str(e)}",
                'response_time': response_time,
                'model_used': f"{self.model_name}",
                'success': False,
                'error': str(e)
            }


class RAGSystem:
    """RAG系统主类"""
    
    def __init__(self):
        self.document_processor = DocumentProcessor()
        self.text_splitter = TextSplitter()
        self.knowledge_bases = {}  # 存储每个知识库的向量存储
        self.llm_configs = {}  # 存储LLM配置
        
    def get_or_create_vector_store(self, kb_id: int) -> VectorStore:
        """获取或创建知识库的向量存储"""
        if kb_id not in self.knowledge_bases:
            self.knowledge_bases[kb_id] = VectorStore()
            # 注意：不在这里自动加载文档，由ask_question方法控制加载时机
        return self.knowledge_bases[kb_id]
    
    def _load_existing_documents(self, kb_id: int):
        """从数据库加载已有的文档数据到向量存储"""
        try:
            from apps.knowledge.models import DocumentChunk
            import asyncio
            
            # 检查是否在异步环境中
            try:
                loop = asyncio.get_event_loop()
                if loop.is_running():
                    # 在异步环境中，使用同步方式强制加载
                    logger.info(f"在异步环境中强制加载知识库 {kb_id} 的数据")
                    # 创建新的线程来执行同步数据库操作
                    import threading
                    import queue
                    
                    result_queue = queue.Queue()
                    
                    def sync_load():
                        try:
                            # 获取知识库中所有已完成的文档块
                            chunks = DocumentChunk.objects.filter(
                                document__knowledge_base_id=kb_id,
                                document__status='completed'
                            ).select_related('document')
                            
                            chunk_data = []
                            for chunk in chunks:
                                chunk_data.append({
                                    'content': chunk.content,
                                    'metadata': {
                                        'document_id': chunk.document.id,
                                        'chunk_index': chunk.chunk_index,
                                        'source': chunk.document.file_path,
                                        'type': chunk.document.file_type,
                                        **chunk.metadata
                                    }
                                })
                            result_queue.put(chunk_data)
                        except Exception as e:
                            result_queue.put(e)
                    
                    thread = threading.Thread(target=sync_load)
                    thread.start()
                    thread.join()
                    
                    chunk_data = result_queue.get()
                    if isinstance(chunk_data, Exception):
                        raise chunk_data
                    
                    vector_store = self.knowledge_bases[kb_id]
                    for chunk_info in chunk_data:
                        vector_store.chunks.append(chunk_info['content'])
                        vector_store.metadata.append(chunk_info['metadata'])
                    
                    # 重建向量
                    if vector_store.chunks:
                        vector_store._update_vectors()
                    
                    logger.info(f"从数据库加载知识库 {kb_id} 的文档数据: {len(vector_store.chunks)} 个块")
                    return
                    
            except RuntimeError:
                # 没有事件循环，可以进行同步操作
                pass
            
            # 同步环境中的正常加载
            # 获取知识库中所有已完成的文档块
            chunks = DocumentChunk.objects.filter(
                document__knowledge_base_id=kb_id,
                document__status='completed'
            ).select_related('document')
            
            vector_store = self.knowledge_bases[kb_id]
            
            for chunk in chunks:
                # 添加文档内容和元数据
                vector_store.chunks.append(chunk.content)
                vector_store.metadata.append({
                    'document_id': chunk.document.id,
                    'chunk_index': chunk.chunk_index,
                    'source': chunk.document.file_path,
                    'type': chunk.document.file_type,
                    **chunk.metadata
                })
            
            # 重建向量
            if vector_store.chunks:
                vector_store._update_vectors()
                
            logger.info(f"从数据库加载知识库 {kb_id} 的文档数据: {len(vector_store.chunks)} 个块")
            
        except Exception as e:
            logger.error(f"加载知识库 {kb_id} 的文档数据失败: {e}")
            import traceback
            traceback.print_exc()
    
    def manually_load_documents(self, kb_id: int):
        """手动加载知识库文档数据（同步方法）- 确保每次都能成功加载"""
        try:
            from apps.knowledge.models import DocumentChunk, Document
            
            logger.info(f"开始强制加载知识库 {kb_id} 的文档数据")
            
            # 首先检查数据库中是否有文档 - 使用线程安全的方式
            doc_count = 0
            try:
                import threading
                count_holder = [0]
                exception_holder = [None]
                
                def count_docs():
                    try:
                        count = Document.objects.filter(
                            knowledge_base_id=kb_id,
                            status='completed'
                        ).count()
                        count_holder[0] = count
                    except Exception as e:
                        exception_holder[0] = e
                
                thread = threading.Thread(target=count_docs)
                thread.start()
                thread.join()
                
                if exception_holder[0]:
                    raise exception_holder[0]
                    
                doc_count = count_holder[0]
                
            except Exception as e:
                logger.error(f"查询文档数量失败: {e}")
                return 0
            
            logger.info(f"数据库中知识库 {kb_id} 有 {doc_count} 个已完成文档")
            
            if doc_count == 0:
                logger.info(f"知识库 {kb_id} 中没有已完成的文档")
                return 0
            
            # 确保向量存储已初始化
            if kb_id not in self.knowledge_bases:
                self.knowledge_bases[kb_id] = VectorStore()
                logger.info(f"为知识库 {kb_id} 创建新的向量存储")
            
            vector_store = self.knowledge_bases[kb_id]
            
            # 强制清空现有数据，重新加载以确保数据同步
            vector_store.chunks.clear()
            vector_store.metadata.clear()
            logger.info(f"清空知识库 {kb_id} 的现有向量数据")
            
            # 获取知识库中所有已完成的文档块 - 简化查询并包装为线程安全
            try:
                # 在新线程中执行数据库查询以避免异步上下文问题
                import threading
                chunks_list = []
                exception_holder = [None]
                
                def fetch_chunks():
                    try:
                        chunks = DocumentChunk.objects.filter(
                            document__knowledge_base_id=kb_id,
                            document__status='completed'
                        ).select_related('document').order_by('id')
                        chunks_list.extend(list(chunks))
                    except Exception as e:
                        exception_holder[0] = e
                
                thread = threading.Thread(target=fetch_chunks)
                thread.start()
                thread.join()
                
                if exception_holder[0]:
                    raise exception_holder[0]
                    
                chunks = chunks_list
                
            except Exception as e:
                logger.error(f"数据库查询失败: {e}")
                return 0
            
            chunk_count = 0
            for chunk in chunks:
                try:
                    # 添加文档内容和元数据
                    vector_store.chunks.append(chunk.content)
                    vector_store.metadata.append({
                        'document_id': chunk.document.id,
                        'chunk_index': chunk.chunk_index,
                        'source': chunk.document.file_path,
                        'type': chunk.document.file_type,
                        **chunk.metadata
                    })
                    chunk_count += 1
                    
                    # 每处理10个chunk记录一次日志
                    if chunk_count % 10 == 0:
                        logger.info(f"已处理 {chunk_count} 个文档块")
                        
                except Exception as e:
                    logger.error(f"处理文档块 {chunk.id} 时出错: {e}")
                    continue
            
            logger.info(f"总共处理了 {chunk_count} 个文档块")
            
            # 验证数据加载
            if vector_store.chunks:
                logger.info(f"向量存储中现在有 {len(vector_store.chunks)} 个块")
                logger.info(f"第一个块内容预览: {vector_store.chunks[0][:100]}...")
                
                # 重建向量
                try:
                    vector_store._update_vectors()
                    logger.info(f"知识库 {kb_id} 成功重建向量索引")
                except Exception as e:
                    logger.error(f"重建向量索引失败: {e}")
            else:
                logger.error(f"知识库 {kb_id} 加载后仍然没有任何数据块！")
            
            logger.info(f"成功强制加载知识库 {kb_id} 的文档数据: {chunk_count} 个块")
            return chunk_count
            
        except Exception as e:
            logger.error(f"强制加载知识库 {kb_id} 的文档数据失败: {e}")
            import traceback
            traceback.print_exc()
            return 0

    def configure_llm(self, config_id: int, model_config: Dict):
        """配置大语言模型"""
        self.llm_configs[config_id] = LLMInterface(model_config)
    
    def process_document(self, kb_id: int, file_path: str, document_id: Optional[int] = None) -> Dict:
        """处理文档"""
        try:
            # 处理文档
            content, metadata = self.document_processor.process_file(file_path)
            
            # 为每个块添加document_id到元数据中
            if document_id:
                metadata['document_id'] = document_id
            
            # 分块
            chunks = self.text_splitter.split_text(content, metadata)
            
            # 获取向量存储并添加文档
            vector_store = self.get_or_create_vector_store(kb_id)
            vector_store.add_documents(chunks)
            
            return {
                'success': True,
                'chunk_count': len(chunks),
                'content_length': len(content),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"处理文档失败: {e}")
            return {
                'success': False,
                'error': str(e),
                'chunk_count': 0
            }
    
    async def ask_question(self, kb_id: int, question: str, config_id: Optional[int] = None, 
                          top_k: int = 5, threshold: float = 0.5) -> Dict:
        """智能问答"""
        import time
        start_time = time.time()
        
        try:
            # 获取向量存储
            vector_store = self.get_or_create_vector_store(kb_id)
            
            # 每次问答都重新检查并加载文档，确保文档数据是最新的
            logger.info(f"知识库 {kb_id} 开始检查文档状态")
            
            # 强制重新加载文档数据以确保数据完整性
            loaded_count = self.manually_load_documents(kb_id)
            logger.info(f"知识库 {kb_id} 重新加载了 {loaded_count} 个文档块")
            
            # 重新获取向量存储（确保获取最新数据）
            vector_store = self.get_or_create_vector_store(kb_id)
            
            # 检索相关文档 - 使用更低的阈值确保能检索到文档
            relevant_docs = vector_store.similarity_search(question, top_k=top_k, threshold=max(threshold, 0.1))
            
            logger.info(f"检索到 {len(relevant_docs)} 个相关文档片段，阈值: {max(threshold, 0.1)}")
            
            # 如果没有检索到文档，尝试降低阈值再次检索
            if not relevant_docs and threshold > 0.0:
                logger.info("未找到相关文档，尝试降低阈值重新检索")
                relevant_docs = vector_store.similarity_search(question, top_k=top_k, threshold=0.0)
                logger.info(f"降低阈值后检索到 {len(relevant_docs)} 个文档片段")
            
            # 构建上下文 - 强制使用知识库内容，确保总是有内容
            context = ""
            context_info = ""
            
            # 首先尝试使用检索到的相关文档
            if relevant_docs:
                context = "\n".join([doc['content'] for doc in relevant_docs])
                context_info = f"基于知识库中的 {len(relevant_docs)} 个相关文档片段："
                logger.info(f"使用相关文档构建上下文，长度: {len(context)} 字符")
            
            # 如果没有相关文档但有知识库内容，强制使用前几个块
            if not context and vector_store.chunks:
                logger.info("没有找到相关文档，强制使用知识库前几个文档块")
                context = "\n".join(vector_store.chunks[:min(10, len(vector_store.chunks))])
                context_info = f"基于知识库中的前 {min(10, len(vector_store.chunks))} 个文档片段："
                logger.info(f"强制构建的上下文长度: {len(context)} 字符")
                
                # 同时将前几个块当作relevant_docs处理，保证后续逻辑正确
                relevant_docs = []
                for i, chunk in enumerate(vector_store.chunks[:min(10, len(vector_store.chunks))]):
                    relevant_docs.append({
                        'content': chunk,
                        'score': 0.1,  # 给一个默认分数
                        'metadata': vector_store.metadata[i] if i < len(vector_store.metadata) else {},
                        'index': i
                    })
            
            # 最后的保险：如果仍然没有context，检查是否真的没有数据
            if not context:
                # 再次尝试直接从数据库获取一些内容
                try:
                    from apps.knowledge.models import DocumentChunk
                    import threading
                    
                    db_content = []
                    db_chunks_data = []
                    exception_holder = [None]
                    
                    def fetch_db_chunks():
                        try:
                            db_chunks = DocumentChunk.objects.filter(
                                document__knowledge_base_id=kb_id,
                                document__status='completed'
                            ).select_related('document')[:5]
                            
                            for chunk in db_chunks:
                                db_content.append(chunk.content)
                                db_chunks_data.append(chunk)
                        except Exception as e:
                            exception_holder[0] = e
                    
                    thread = threading.Thread(target=fetch_db_chunks)
                    thread.start()
                    thread.join()
                    
                    if exception_holder[0]:
                        raise exception_holder[0]
                    
                    if db_content:
                        logger.warning("向量存储为空但数据库有数据，直接从数据库获取")
                        context = "\n".join(db_content)
                        context_info = f"直接从数据库获取的 {len(db_content)} 个文档片段："
                        logger.info(f"从数据库直接获取的上下文长度: {len(context)} 字符")
                        
                        # 构造相应的relevant_docs
                        relevant_docs = []
                        for i, chunk in enumerate(db_chunks_data):
                            relevant_docs.append({
                                'content': chunk.content,
                                'score': 0.05,  # 更低的分数表示这是直接获取的
                                'metadata': {'document_id': chunk.document.id, 'chunk_index': chunk.chunk_index},
                                'index': i
                            })
                    else:
                        context = ""
                        context_info = "知识库中没有找到任何文档"
                        logger.error(f"知识库 {kb_id} 数据库中也没有任何文档内容")
                except Exception as e:
                    logger.error(f"从数据库获取备用内容失败: {e}")
                    context = ""
                    context_info = "知识库读取失败"
            
            # 记录最终的context状态
            logger.info(f"最终context状态: 长度={len(context)}, 信息={context_info}")
            if context:
                logger.info(f"Context前200字符: {context[:200]}...")
            
            # 生成回答 - 确保总是将知识库内容传递给大模型
            llm = self.llm_configs.get(config_id) if config_id else None
            if llm:
                logger.info(f"使用LLM配置ID: {config_id}")
                logger.info(f"检查上下文状态: context长度={len(context) if context else 0}, vector_store.chunks数量={len(vector_store.chunks)}, relevant_docs数量={len(relevant_docs)}")
                
                # 最后的强制保险：如果context仍然为空，直接从数据库强制获取
                if not context:
                    logger.error("严重警告: context为空，执行最终兜底操作")
                    try:
                        from apps.knowledge.models import DocumentChunk
                        import threading
                        
                        emergency_content = []
                        exception_holder = [None]
                        
                        def fetch_emergency_chunks():
                            try:
                                emergency_chunks = DocumentChunk.objects.filter(
                                    document__knowledge_base_id=kb_id,
                                    document__status='completed'
                                )[:3]
                                
                                for chunk in emergency_chunks:
                                    emergency_content.append(chunk.content)
                            except Exception as e:
                                exception_holder[0] = e
                        
                        thread = threading.Thread(target=fetch_emergency_chunks)
                        thread.start()
                        thread.join()
                        
                        if exception_holder[0]:
                            raise exception_holder[0]
                        
                        if emergency_content:
                            context = "\n".join(emergency_content)
                            logger.error(f"紧急兜底: 从数据库获取到 {len(emergency_content)} 个块")
                    except Exception as emergency_e:
                        logger.error(f"紧急兜底也失败: {emergency_e}")
                
                # 现在context应该总是有内容（除非知识库真的为空）
                if context:
                    logger.info(f"使用有内容的context构建提示词，context前100字符: {context[:100]}")
                    # 构建极其明确的提示，强制大模型按格式回答
                    enhanced_question = f"""【严格指令 - 必须遵守】你是专业知识库助手，必须严格按照以下格式回答，不得违反：

🔴 强制要求：
1. 第一句话必须是："基于知识库内容，我为您回答："
2. 禁止使用"基于通用知识"等其他开头
3. 必须引用下面的知识库内容
4. 不得说"没有相关内容"

📚 知识库内容：
{context}

❓ 用户问题：{question}

✅ 回答格式示例：
基于知识库内容，我为您回答：[根据上述知识库内容的具体回答]

⚠️ 重要提醒：无论如何都必须以"基于知识库内容，我为您回答："开头，这是不可违反的规则！

现在请严格按照格式开始回答："""
                else:
                    # 这种情况现在应该极少发生
                    logger.error("即使经过所有兜底措施，context仍然为空！这不应该发生。")
                    enhanced_question = f"""【严格指令】知识库助手必须回答：

第一句话必须是："基于知识库内容，我为您回答："
然后说明："当前知识库系统出现问题，无法读取文档内容。"

用户问题：{question}

请严格按照上述格式回答。"""
                
                logger.info(f"发送给大模型的完整提示长度: {len(enhanced_question)} 字符")
                logger.info(f"上下文内容预览: {context[:300]}..." if context else "上下文为空")
                
                # 使用构建好的完整提示词
                llm_result = await llm.generate_response(enhanced_question, "")
                answer = llm_result.get('answer', '生成回答失败')
                model_used = llm_result.get('model_used', f"config_{config_id}" if config_id else "default")
            else:
                logger.warning(f"没有找到LLM配置，配置ID: {config_id}")
                # 如果没有配置LLM，才返回"未找到相关信息"的提示
                if not relevant_docs:
                    response_time = round((time.time() - start_time), 3)
                    return {
                        'answer': '抱歉，我在知识库中没有找到相关信息，且未配置大语言模型。请配置模型以获得智能回答。',
                        'sources': [],
                        'confidence': 0.0,
                        'retrieved_chunks': [],
                        'model_used': f"config_{config_id}" if config_id else "default",
                        'response_time': response_time
                    }
                else:
                    answer = f"基于知识库内容，找到了 {len(relevant_docs)} 个相关片段，但未配置大语言模型。请配置模型以获得智能回答。"
                    model_used = f"config_{config_id}" if config_id else "default"
            
            response_time = round((time.time() - start_time), 3)  # 保持为秒，保留3位小数
            
            logger.info(f"问答完成: 回答长度={len(answer)}, 源文档数={len(relevant_docs)}, 使用模型={model_used}, 响应时间={response_time}秒")
            
            return {
                'answer': answer,
                'sources': [
                    {
                        'content': doc['content'][:200] + '...' if len(doc['content']) > 200 else doc['content'],
                        'score': doc['score'],
                        'metadata': doc['metadata']
                    }
                    for doc in relevant_docs
                ],
                'confidence': relevant_docs[0]['score'] if relevant_docs else 0.0,
                'retrieved_chunks': relevant_docs,
                'model_used': model_used,
                'response_time': response_time
            }
            
        except Exception as e:
            logger.error(f"问答失败: {e}")
            response_time = round((time.time() - start_time), 3)  # 保持为秒，保留3位小数
            return {
                'answer': f'查询过程中出现错误: {str(e)}',
                'sources': [],
                'confidence': 0.0,
                'retrieved_chunks': [],
                'model_used': "error",
                'response_time': response_time
            }
    
    def get_knowledge_base_stats(self, kb_id: int) -> Dict:
        """获取知识库统计信息"""
        if kb_id in self.knowledge_bases:
            vector_store = self.knowledge_bases[kb_id]
            return {
                'total_chunks': len(vector_store.chunks),
                'total_documents': len(set(chunk.get('document_id', 0) for chunk in vector_store.metadata)),
                'vector_dimension': vector_store.vectors.shape[1] if vector_store.vectors is not None else 0
            }
        else:
            return {
                'total_chunks': 0,
                'total_documents': 0,
                'vector_dimension': 0
            }
